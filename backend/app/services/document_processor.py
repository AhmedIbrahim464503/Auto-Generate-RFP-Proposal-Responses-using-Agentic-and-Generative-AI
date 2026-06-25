import os
import fitz  # PyMuPDF
import docx
from abc import ABC, abstractmethod
from typing import Dict, Any, Type
from backend.app.core.logger import logger

class BaseDocumentProcessor(ABC):
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """
        Extracts raw text content from the file.
        """
        pass

    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts document structure and metadata (page/word counts).
        """
        pass


class PDFProcessor(BaseDocumentProcessor):
    def extract_text(self, file_path: str) -> str:
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise e

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = fitz.open(file_path)
            page_count = doc.page_count
            metadata = doc.metadata
            doc.close()
            
            # Simple word count approximation
            text_content = self.extract_text(file_path)
            word_count = len(text_content.split())

            return {
                "page_count": page_count,
                "word_count": word_count,
                "author": metadata.get("author", "Unknown"),
                "creator": metadata.get("creator", "Unknown"),
                "title": metadata.get("title", "Unknown"),
                "format": "PDF"
            }
        except Exception as e:
            logger.error(f"Error extracting metadata from PDF {file_path}: {str(e)}")
            raise e


class DOCXProcessor(BaseDocumentProcessor):
    def extract_text(self, file_path: str) -> str:
        text = []
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text.append(para.text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise e

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = docx.Document(file_path)
            # Aproxiate structure metrics
            text_content = self.extract_text(file_path)
            word_count = len(text_content.split())
            paragraph_count = len(doc.paragraphs)

            # Metadata properties
            props = doc.core_properties
            
            return {
                "page_count": 1,  # Word doesn't store explicit page counts easily without rendering engine
                "word_count": word_count,
                "paragraph_count": paragraph_count,
                "author": props.author or "Unknown",
                "title": props.title or "Unknown",
                "format": "DOCX"
            }
        except Exception as e:
            logger.error(f"Error extracting metadata from DOCX {file_path}: {str(e)}")
            raise e


class DocumentProcessorFactory:
    _processors: Dict[str, Type[BaseDocumentProcessor]] = {
        ".pdf": PDFProcessor,
        ".docx": DOCXProcessor
    }

    @classmethod
    def get_processor(cls, file_path: str) -> BaseDocumentProcessor:
        _, ext = os.path.splitext(file_path.lower())
        processor_class = cls._processors.get(ext)
        if not processor_class:
            raise ValueError(f"Unsupported file format for processing: {ext}")
        return processor_class()
