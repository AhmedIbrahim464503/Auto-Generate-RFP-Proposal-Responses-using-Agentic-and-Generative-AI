import os
import io
from docx import Document
from sqlalchemy.orm import Session
from backend.app.models.proposal import ProposalPlan, ProposalSection

class ExporterService:
  @staticmethod
  def compile_to_docx(db: Session, proposal_id: str) -> io.BytesIO:
    proposal = db.query(ProposalPlan).filter(ProposalPlan.id == proposal_id).first()
    if not proposal:
      raise ValueError("Proposal not found")

    doc = Document()
    doc.add_heading(proposal.title or "Proposal Document", 0)

    from backend.app.models.proposal_generation import GeneratedSection

    sections = db.query(ProposalSection).filter(ProposalSection.proposal_plan_id == proposal_id).all()
    for sec in sections:
      doc.add_heading(sec.title or "Untitled Section", level=1)
      
      # Look up generated content for this specific section
      gen_sec = db.query(GeneratedSection).filter(
          GeneratedSection.proposal_section_id == sec.id
      ).order_by(GeneratedSection.created_at.desc()).first()
      
      content = gen_sec.content if gen_sec else (sec.content or "Draft content pending generation.")
      doc.add_paragraph(content)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

  @staticmethod
  def compile_to_markdown(db: Session, proposal_id: str) -> str:
    from backend.app.models.proposal_generation import GeneratedSection

    proposal = db.query(ProposalPlan).filter(ProposalPlan.id == proposal_id).first()
    if not proposal:
      raise ValueError("Proposal not found")

    md = []
    md.append(f"# {proposal.title or 'Proposal Document'}")
    md.append("")

    sections = db.query(ProposalSection).filter(ProposalSection.proposal_plan_id == proposal_id).all()
    for sec in sections:
      md.append(f"## {sec.title or 'Untitled Section'}")
      md.append("")
      
      # Look up generated content for this specific section
      gen_sec = db.query(GeneratedSection).filter(
          GeneratedSection.proposal_section_id == sec.id
      ).order_by(GeneratedSection.created_at.desc()).first()
      
      content = gen_sec.content if gen_sec else (sec.content or "Draft content pending generation.")
      md.append(content)
      md.append("")

    return "\n".join(md)
